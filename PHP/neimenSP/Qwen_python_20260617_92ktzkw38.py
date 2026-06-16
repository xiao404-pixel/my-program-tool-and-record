from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 建立文件
doc = Document()

# 設定基本字體
style = doc.styles['Normal']
font = style.font
font.name = '微軟正黑體'
font.size = Pt(12)

# --- 封面/標題 ---
title = doc.add_heading('動態網頁設計專案報告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph('專案名稱：內門區景點與特有活動資料管理系統')
doc.add_paragraph('課程名稱：動態網頁設計')
doc.add_paragraph('學生姓名：_______________')
doc.add_paragraph('學號：_______________')
doc.add_paragraph('提交日期：2025年 6月 12日')

doc.add_page_break()

# --- 目錄 ---
doc.add_heading('目錄', level=1)
toc_items = [
    '一、前言',
    '二、系統開發環境與工具',
    '三、系統功能與畫面展示',
    '四、資料庫設計',
    '五、實作過程與問題排除',
    '六、學習心得與結語'
]
for item in toc_items:
    doc.add_paragraph(item)

doc.add_page_break()

# --- 一、前言 ---
doc.add_heading('一、前言', level=1)
doc.add_paragraph('本次專案的目的是為了建立一個「內門區景點與特有活動資料管理系統」。內門區擁有豐富的觀光資源與特有活動（如觀音媽廟、飛爐傳奇等），為了讓這些資訊能夠被有效管理與更新，我設計了這個基於 PHP 與 MySQL 的網頁後台管理系統。')
doc.add_paragraph('系統主要提供管理員進行景點資料的新增、查詢、修改與刪除（CRUD），並支援圖片上傳與 Google Maps 地圖連結，希望能提供一個簡單且實用的資料管理介面。')

# --- 二、系統開發環境與工具 ---
doc.add_heading('二、系統開發環境與工具', level=1)
doc.add_paragraph('本系統採用常見的 LAMP/WAMP 架構進行開發，具體使用的環境與工具如下：')
doc.add_paragraph('• 伺服器環境：XAMPP (Apache + MySQL + PHP)')
doc.add_paragraph('• 後端語言：PHP 7.x (使用 PDO 進行資料庫連線)')
doc.add_paragraph('• 資料庫：MySQL (資料庫名稱：neimensp)')
doc.add_paragraph('• 前端技術：HTML5, CSS3, JavaScript')
doc.add_paragraph('• 程式編輯器：Visual Studio Code')

# --- 三、系統功能與畫面展示 ---
doc.add_heading('三、系統功能與畫面展示', level=1)

doc.add_heading('3.1 會員登入 (login.php)', level=2)
doc.add_paragraph('系統設有簡單的帳號密碼驗證機制。管理者輸入正確的帳號密碼後，系統會利用 Session 記錄登入狀態，並跳轉至首頁。若未登入則無法直接存取後台頁面。')
doc.add_paragraph('【畫面截圖】')
doc.add_paragraph('（請在此處插入登入頁面截圖）')

doc.add_heading('3.2 首頁與資料列表 (home.php)', level=2)
doc.add_paragraph('登入後的首頁會以表格形式顯示所有景點資料。為了避免資料過多，我實作了分頁功能，每頁顯示 3 筆資料。此外，表格中的「景點活動編號」是透過 PHP 計數器動態生成的，這樣即使刪除了中間的資料，編號依然會自動順排，不會出現斷號。')
doc.add_paragraph('【畫面截圖】')
doc.add_paragraph('（請在此處插入首頁資料列表截圖）')

doc.add_heading('3.3 新增景點 (add.php)', level=2)
doc.add_paragraph('管理者可以透過表單新增景點資訊。除了基本的文字欄位外，也支援圖片上傳（會自動產生唯一檔名避免覆蓋），以及選填的 Google Maps 地圖導航連結。送出前會有 JavaScript 確認視窗防止誤觸。')
doc.add_paragraph('【畫面截圖】')
doc.add_paragraph('（請在此處插入新增頁面截圖）')

doc.add_heading('3.4 修改景點 (update.php)', level=2)
doc.add_paragraph('修改頁面會預先帶入該筆景點的現有資料。在圖片處理上，如果管理者沒有選擇新圖片，系統會自動保留原本的舊圖片，不會導致破圖。若上傳新圖，則會自動刪除伺服器上的舊圖片檔。')
doc.add_paragraph('【畫面截圖】')
doc.add_paragraph('（請在此處插入修改頁面截圖）')

doc.add_heading('3.5 刪除景點 (delete.php)', level=2)
doc.add_paragraph('點擊刪除後，會先進入確認頁面，顯示即將刪除的景點名稱與圖片。確認無誤並點擊按鈕後，會觸發 JavaScript 確認視窗，通過後才會真正從資料庫刪除資料，並同步清除伺服器上的圖片檔案。')
doc.add_paragraph('【畫面截圖】')
doc.add_paragraph('（請在此處插入刪除確認頁面截圖）')

doc.add_page_break()

# --- 四、資料庫設計 ---
doc.add_heading('四、資料庫設計', level=1)
doc.add_paragraph('本系統使用 MySQL 資料庫，主要包含以下兩個資料表：')

doc.add_heading('4.1 景點資料表 (products)', level=2)
table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
headers = ['欄位名稱', '資料類型', '說明']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

fields = [
    ('pno', 'INT', '景點編號 (Primary Key, Auto Increment)'),
    ('pname', 'VARCHAR(255)', '景點活動名稱'),
    ('pptime', 'VARCHAR(255)', '景點開放時間'),
    ('pscrib', 'TEXT', '景點詳細描述'),
    ('ppic', 'VARCHAR(255)', '圖片檔名'),
    ('pmap_link', 'VARCHAR(500)', '地圖導航連結 (Google Maps URL)')
]
for i, (field, dtype, desc) in enumerate(fields, 1):
    table.rows[i].cells[0].text = field
    table.rows[i].cells[1].text = dtype
    table.rows[i].cells[2].text = desc

doc.add_heading('4.2 使用者資料表 (users)', level=2)
doc.add_paragraph('用於儲存管理員的帳號與密碼，包含 id, username, pwd 三個欄位。')

# --- 五、實作過程與問題排除 ---
doc.add_heading('五、實作過程與問題排除', level=1)
doc.add_paragraph('在開發過程中，我遇到了一些問題，但透過查閱資料與除錯，最終都成功解決了：')

doc.add_heading('問題一：修改資料時，若未上傳新圖片會導致破圖', level=2)
doc.add_paragraph('【問題描述】：剛開始寫 update.php 時，如果管理者只修改文字而沒有選擇新圖片，送出後資料庫的圖片欄位會變成空白，導致首頁圖片破圖。')
doc.add_paragraph('【解決方法】：我在處理表單資料時，先將舊圖片的檔名 (`old_ppic`) 賦值給 `$ppic` 變數。只有當偵測到有上傳新檔案且成功時，才會將 `$ppic` 覆蓋為新檔名，並使用 `unlink()` 刪除舊圖。這樣就能完美保留舊圖片。')

doc.add_heading('問題二：新增地圖連結後，首頁依然顯示「無」', level=2)
doc.add_paragraph('【問題描述】：我在表單中加入了地圖連結的輸入框，但儲存後首頁卻讀不到資料。')
doc.add_paragraph('【解決方法】：檢查後發現，我雖然有接收 `$_POST[\'pmap_link\']`，但忘記在 INSERT 和 UPDATE 的 SQL 語句中加上 `pmap_link` 欄位，首頁的 SELECT 語句也漏了這個欄位。補齊 SQL 語句後問題就解決了。')

doc.add_heading('問題三：刪除資料後，編號出現斷號', level=2)
doc.add_paragraph('【問題描述】：原本直接使用資料庫的 `pno` 作為顯示編號，刪除第 2 筆後，編號會變成 1, 3, 4... 看起來不連貫。')
doc.add_paragraph('【解決方法】：我改變了顯示邏輯，不使用資料庫的 `pno`，而是在 `home.php` 使用 `$counter = $offset + 1` 來計算當前頁面的顯示序號，這樣無論資料庫怎麼刪除，前端顯示的編號永遠是連續的 1, 2, 3...')

# --- 六、學習心得與結語 ---
doc.add_heading('六、學習心得與結語', level=1)
doc.add_paragraph('透過這次「動態網頁設計」的專案實作，我將課堂上學到的 PHP 基礎語法、MySQL 資料庫操作以及前後端互動的知識，實際應用到了一個完整的系統中。')
doc.add_paragraph('我學會了如何使用 PDO 來安全地連接資料庫並防止 SQL 注入，也掌握了 Session 登入驗證的機制。在實作檔案上傳與圖片管理的過程中，更讓我了解到伺服器檔案處理的細節（如權限、檔名衝突、舊檔刪除等）。')
doc.add_paragraph('雖然這個系統目前還比較陽春，未來如果有機會，我還想加入搜尋功能、分類管理，以及將密碼進行雜湊加密（如 bcrypt）來提升系統的安全性。這次專案讓我對網頁後端開發有了更深刻的認識，也累積了寶貴的實作經驗。')

# 儲存文件
doc.save('動態網頁設計專案報告.docx')
print("✅ 報告書已成功產生：動態網頁設計專案報告.docx")